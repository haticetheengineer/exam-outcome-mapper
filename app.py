import streamlit as st
import json
import re
import io
import requests
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

st.set_page_config(
    page_title="Exam Outcome Mapper",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── SESSION STATE ─────────────────────────────────────────────
for k, v in {
    "logged_in": False,
    "lang": "TR",
    "dark": False,
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── DİL ──────────────────────────────────────────────────────
TR = {
    "title": "Sınav ÖÇ Eşleştirme",
    "subtitle": "Kapadokya Üniversitesi — Akreditasyon Rapor Sistemi",
    "login": "Giriş Yap",
    "username": "Kullanıcı Adı",
    "password": "Şifre",
    "username_ph": "Kullanıcı adınızı girin",
    "password_ph": "Şifrenizi girin",
    "login_btn": "Giriş Yap",
    "login_err": "❌ Hatalı kullanıcı adı veya şifre",
    "start_over": "🔄 Yeniden Başla",
    "logout": "🚪 Çıkış",
    "step1": "📂 Adım 1 — Sınav Dosyası Yükle",
    "step1_sub": "TXT, DOCX veya PDF",
    "step1_ok": "soru bulundu",
    "step1_err": "❌ Soru bulunamadı. Sorular '?' ile bitmelidir.",
    "preview": "Önizleme",
    "step2": "🎯 Adım 2 — Öğrenim Çıktıları",
    "tab_pdf": "📄 PDF'den Çıkar",
    "tab_manual": "✏️ Elle Gir",
    "pdf_info": "Bilgi Paketi PDF yükleyin — ÖÇ'ler otomatik çıkarılır.",
    "pdf_upload": "Bilgi Paketi PDF",
    "pdf_ok": "ÖÇ çıkarıldı",
    "pdf_err": "❌ ÖÇ bulunamadı",
    "manual_info": "💡 Her satıra: `LO-1: tanım`",
    "manual_ph": "LO-1: Makine öğrenmesi algoritmalarını açıklar\nLO-2: Yöntemleri uygular",
    "outcomes_ok": "öğrenim çıktısı tanımlandı",
    "step3": "💯 Adım 3 — Puanlama",
    "total_pts": "Toplam sınav puanı",
    "scoring_type": "Soru puanlama",
    "equal": "Eşit (otomatik)",
    "custom": "Özel (soru bazlı)",
    "enter_scores": "Her soru için puan girin:",
    "total_warn": "⚠️ Toplam:",
    "total_ok": "✅ Toplam:",
    "answer_key": "📝 Cevap Anahtarı (opsiyonel)",
    "answer_ph": "ABCDEABCDE...",
    "step4": "📊 Adım 4 — Rapor Oluştur",
    "ready_equal": "Hazır:",
    "ready_pts": "puan/soru",
    "ready_custom": "özel puanlama",
    "questions": "soru",
    "outcomes": "çıktı",
    "map_btn": "🚀 AI ile Eşleştir ve Excel Oluştur",
    "mapped_ok": "soru eşleştirildi",
    "multi_match": "birden fazla ÖÇ ile eşleşti",
    "download": "⬇ Excel Raporu İndir (.xlsx)",
    "no_exam": "⬆ Önce sınav dosyası yükleyin",
    "no_oc": "⬆ Önce öğrenim çıktılarını tanımlayın",
    "extracting": "Sorular çıkarılıyor...",
    "oc_extracting": "ÖÇ'ler çıkarılıyor...",
    "mapping": "Eşleştirme yapılıyor...",
    "batch": "Batch",
    "rubrik_info": "Rubrik Form Bilgileri",
    "ogretim_elemani": "Öğretim Elemanı",
    "sinav_turu": "Sınav Türü",
    "bolum": "Bölüm / Program",
    "sinav_tarihi": "Sınav Tarihi",
    "download_rubrik": "Rubrik Form İndir (.docx)",
}

EN = {
    "title": "Exam Outcome Mapper",
    "subtitle": "Kapadokya University — Accreditation Report System",
    "login": "Login",
    "username": "Username",
    "password": "Password",
    "username_ph": "Enter username",
    "password_ph": "Enter password",
    "login_btn": "Login",
    "login_err": "❌ Invalid username or password",
    "start_over": "🔄 Start Over",
    "logout": "🚪 Logout",
    "step1": "📂 Step 1 — Upload Exam File",
    "step1_sub": "TXT, DOCX or PDF",
    "step1_ok": "questions found",
    "step1_err": "❌ No questions found. Questions must end with '?'",
    "preview": "Preview",
    "step2": "🎯 Step 2 — Learning Outcomes",
    "tab_pdf": "📄 Extract from PDF",
    "tab_manual": "✏️ Enter Manually",
    "pdf_info": "Upload Bilgi Paketi PDF — outcomes extracted automatically.",
    "pdf_upload": "Bilgi Paketi PDF",
    "pdf_ok": "outcomes extracted",
    "pdf_err": "❌ No outcomes found",
    "manual_info": "💡 One per line: `LO-1: definition`",
    "manual_ph": "LO-1: Explains ML algorithms\nLO-2: Applies methods",
    "outcomes_ok": "outcomes defined",
    "step3": "💯 Step 3 — Scoring",
    "total_pts": "Total exam points",
    "scoring_type": "Question scoring",
    "equal": "Equal (auto)",
    "custom": "Custom per question",
    "enter_scores": "Enter score for each question:",
    "total_warn": "⚠️ Total:",
    "total_ok": "✅ Total:",
    "answer_key": "📝 Answer Key (optional)",
    "answer_ph": "ABCDEABCDE...",
    "step4": "📊 Step 4 — Generate Report",
    "ready_equal": "Ready:",
    "ready_pts": "pts each",
    "ready_custom": "custom scoring",
    "questions": "questions",
    "outcomes": "outcomes",
    "map_btn": "🚀 Map with AI & Generate Excel",
    "mapped_ok": "questions mapped",
    "multi_match": "matched to multiple outcomes",
    "download": "⬇ Download Excel Report (.xlsx)",
    "no_exam": "⬆ Upload an exam file first",
    "no_oc": "⬆ Define learning outcomes first",
    "extracting": "Extracting questions...",
    "oc_extracting": "Extracting outcomes...",
    "mapping": "Mapping with AI...",
    "batch": "Batch",
    "rubrik_info": "Rubric Form Info",
    "ogretim_elemani": "Instructor",
    "sinav_turu": "Exam Type",
    "bolum": "Department / Program",
    "sinav_tarihi": "Exam Date",
    "download_rubrik": "Download Rubric Form (.docx)",
}

def t(key):
    d = TR if st.session_state.lang == "TR" else EN
    return d.get(key, key)

# ── TEMA CSS ──────────────────────────────────────────────────
def get_css(dark):
    common = """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;600;700&display=swap');
    /* Toggle görünür yap */
    div[data-testid="stToggle"] {
        background: #e9e6ff !important;
        border: 1.5px solid #c4b5fd !important;
        border-radius: 20px !important;
        padding: 4px 10px !important;
        display: inline-flex !important;
        align-items: center !important;
        gap: 6px !important;
    }
    div[data-testid="stToggle"] label {
        color: #4c3bcf !important;
        font-weight: 600 !important;
        font-size: 0.8rem !important;
    }
    div[data-testid="stToggle"] p {
        color: #4c3bcf !important;
        font-weight: 600 !important;
    }
    /* Sidebar yazıları her zaman görünür */
    section[data-testid="stSidebar"] * {
        opacity: 1 !important;
    }
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] small {
        color: inherit !important;
    }
    /* Sadece metin içeriklerine monospace uygula, butonlara değil */
    p, h1, h2, h3, h4, h5, h6, li, span.stMarkdown,
    div[data-testid="stMarkdownContainer"],
    div[data-testid="stMarkdownContainer"] *,
    .stTextInput input, .stTextArea textarea,
    .stNumberInput input, label {
        font-family: 'JetBrains Mono', monospace !important;
    }
    .block-container { padding-top: 0.5rem !important; max-width: 860px !important; padding-left: 2rem !important; padding-right: 2rem !important; }
    /* Streamlit üst header'ı gizle */
    header[data-testid="stHeader"] { display: none !important; }
    #MainMenu { display: none !important; }
    footer { display: none !important; }
    /* File uploader fix */
    div[data-testid="stFileUploader"] label { display: none !important; }
    div[data-testid="stFileUploaderDropzone"] button {
        font-family: sans-serif !important;
        font-size: 0.85rem !important;
    }
    div[data-testid="stFileUploaderDropzone"] button p {
        font-family: sans-serif !important;
    }
    /* Butonlar */
    .stButton > button { border-radius: 8px !important; font-weight: 600 !important; transition: all 0.2s !important; }
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #7c6af7, #6d28d9) !important;
        border: none !important; color: white !important;
        box-shadow: 0 4px 15px rgba(124,106,247,0.35) !important;
    }
    .stButton > button[kind="primary"]:hover { transform: translateY(-2px) !important; }
    .stDownloadButton > button {
        background: linear-gradient(135deg, #7c6af7, #6d28d9) !important;
        color: white !important; border: none !important;
        border-radius: 8px !important; font-weight: 600 !important;
        box-shadow: 0 4px 15px rgba(124,106,247,0.35) !important;
        width: 100% !important; padding: 12px !important;
    }
    /* App header */
    .app-header {
        padding: 28px 32px; border-radius: 16px;
        margin-bottom: 28px; text-align: center;
    }
    .app-header h1 { margin:0; font-size:1.6rem; font-weight:800; letter-spacing:-0.5px; }
    .app-header p  { margin:6px 0 0; font-size:0.82rem; }
    /* Section title */
    .section-title {
        font-size:0.7rem; font-weight:700; letter-spacing:2.5px;
        text-transform:uppercase; margin-bottom:14px;
        display:flex; align-items:center; gap:10px;
    }
    .section-title::after { content:''; flex:1; height:1px; }
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] { border-radius: 10px !important; padding: 4px !important; }
    .stTabs [data-baseweb="tab"] { border-radius: 8px !important; }
    </style>
    """

    if dark:
        theme = """
    <style>
    /* === DARK MODE: tüm arka planları zorla === */
    .stApp, section.main, div[data-testid="stAppViewContainer"],
    div[data-testid="stMain"], div[data-testid="block-container"],
    .block-container, .main .block-container {
        background-color: #0f1117 !important;
        color: #e8eaf6 !important;
    }
    section[data-testid="stSidebar"] > div:first-child {
        background-color: #131620 !important;
    }
    /* Input alanları */
    input, textarea, [data-baseweb="input"] input,
    [data-baseweb="textarea"] textarea,
    div[data-testid="stTextInput"] input,
    div[data-testid="stTextArea"] textarea,
    div[data-testid="stNumberInput"] input {
        background-color: #1e2130 !important;
        color: #e8eaf6 !important;
        border-color: #3d4166 !important;
    }
    /* Label ve yazılar */
    label, p, span, div, h1, h2, h3, h4,
    .stMarkdown p, .stMarkdown span,
    div[data-testid="stMarkdownContainer"] p {
        color: #e8eaf6 !important;
    }
    /* Alert/info kutuları */
    div[data-testid="stAlert"] {
        background-color: #1a1d2e !important;
    }
    /* Expander */
    div[data-testid="stExpander"],
    div[data-testid="stExpander"] summary,
    details, details summary {
        background-color: #1e2130 !important;
        color: #e8eaf6 !important;
    }
    /* File uploader */
    div[data-testid="stFileUploader"] > div {
        background-color: #1e2130 !important;
        border-color: #3d4166 !important;
    }
    /* Tab listesi */
    div[data-baseweb="tab-list"] {
        background-color: #1e2130 !important;
        border-radius: 10px !important;
        padding: 4px !important;
        gap: 4px !important;
    }
    div[data-baseweb="tab"] {
        white-space: nowrap !important;
        min-width: fit-content !important;
        padding: 6px 16px !important;
        border-radius: 8px !important;
        color: #a0aec0 !important;
    }
    div[data-baseweb="tab"][aria-selected="true"] {
        background-color: #2d3250 !important;
        color: #e8eaf6 !important;
        font-weight: 600 !important;
    }
    /* Radio */
    div[data-testid="stRadio"] > div { color: #e8eaf6 !important; }
    /* Toggle */
    div[data-testid="stToggle"] label { color: #e8eaf6 !important; }
    /* Number input */
    div[data-testid="stNumberInput"] div { color: #e8eaf6 !important; }
    /* Divider */
    hr { border-color: #2d3250 !important; }
    /* App header */
    .app-header {
        background: linear-gradient(135deg, #1e1b4b 0%, #312e81 60%, #4c1d95 100%) !important;
        box-shadow: 0 8px 32px rgba(124,106,247,0.25) !important;
        border: 1px solid rgba(124,106,247,0.2) !important;
    }
    .app-header h1 { color: #e0e7ff !important; }
    .app-header p  { color: #a5b4fc !important; }
    .section-title { color: #a78bfa !important; }
    .section-title::after { background: #2d3250 !important; }
    </style>
    """
    else:
        theme = """
    <style>
    /* Light mode — tüm arka planları zorla */
    .stApp, section.main,
    div[data-testid="stAppViewContainer"],
    div[data-testid="stMain"],
    div[data-testid="block-container"],
    .block-container {
        background-color: #f5f4ff !important;
        color: #1a1640 !important;
    }
    section[data-testid="stSidebar"],
    section[data-testid="stSidebar"] > div:first-child {
        background-color: #ffffff !important;
    }
    section[data-testid="stSidebar"] *,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] small {
        color: #1a1640 !important;
        opacity: 1 !important;
    }
    /* Input alanları */
    input, textarea, div[data-testid="stTextInput"] input,
    div[data-testid="stTextArea"] textarea,
    div[data-testid="stNumberInput"] input {
        background-color: #ffffff !important;
        color: #1a1640 !important;
        border-color: #d4d0f0 !important;
    }
    /* Yazılar */
    p, h1, h2, h3, label, span, div {
        color: #1a1640 !important;
    }
    /* Header */
    .app-header {
        background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 60%, #9333ea 100%);
        box-shadow: 0 8px 32px rgba(109,92,231,0.28);
    }
    .app-header h1 { color: white !important; }
    .app-header p  { color: rgba(255,255,255,0.8) !important; }
    .section-title { color: #6d5ce7 !important; }
    .section-title::after { background: #e2e0f0; }
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        background: #ede9fe !important;
        border-radius: 10px !important;
        padding: 4px !important;
    }
    .stTabs [data-baseweb="tab"] {
        white-space: nowrap !important;
        padding: 8px 18px !important;
        border-radius: 8px !important;
        color: #6e6c8e !important;
    }
    .stTabs [aria-selected="true"] {
        background: white !important;
        color: #6d5ce7 !important;
        font-weight: 600 !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08) !important;
    }
    /* Alert */
    div[data-testid="stAlert"] {
        background-color: #f0edff !important;
        color: #1a1640 !important;
    }
    /* Expander */
    details, details summary {
        background-color: #ffffff !important;
        color: #1a1640 !important;
    }
    /* File uploader - light mode */
    div[data-testid="stFileUploader"],
    div[data-testid="stFileUploader"] > div,
    div[data-testid="stFileUploader"] section,
    div[data-testid="stFileUploaderDropzone"],
    div[data-testid="stFileUploaderDropzone"] > div {
        background-color: #ffffff !important;
        background: #ffffff !important;
        border: 1.5px dashed #c4b5fd !important;
        border-radius: 10px !important;
        color: #1a1640 !important;
    }
    div[data-testid="stFileUploaderDropzone"] button,
    div[data-testid="stFileUploaderDropzone"] button *,
    [data-testid="stFileUploaderDropzone"] button,
    [class*="uploadButton"],
    button[data-testid="baseButton-secondary"] {
        background-color: #7c6af7 !important;
        background: #7c6af7 !important;
        color: white !important;
        border-radius: 8px !important;
        font-family: sans-serif !important;
        border: none !important;
        opacity: 1 !important;
    }
    /* Tüm secondary butonları light modda mor yap */
    .stApp button[kind="secondary"],
    .stApp .stButton button[kind="secondary"] {
        background-color: #7c6af7 !important;
        color: white !important;
        border: none !important;
    }
    div[data-testid="stFileUploaderDropzone"] small,
    div[data-testid="stFileUploaderDropzone"] span:not(button span),
    div[data-testid="stFileUploaderDropzone"] p {
        color: #6e6c8e !important;
    }
    /* Sidebar butonları - light mode */
    section[data-testid="stSidebar"] .stButton > button {
        background: #f0edff !important;
        color: #4c3bcf !important;
        border: 1px solid #c4b5fd !important;
        font-weight: 600 !important;
    }
    section[data-testid="stSidebar"] .stButton > button:hover {
        background: #7c6af7 !important;
        color: white !important;
        border-color: #7c6af7 !important;
    }
    </style>
    """
    return common + theme
st.markdown(get_css(st.session_state.dark), unsafe_allow_html=True)

# ── LOGIN ─────────────────────────────────────────────────────
if not st.session_state.logged_in:
    st.markdown(f"""<div class='app-header'>
        <h1>📋 {t('title')}</h1>
        <p>{t('subtitle')}</p>
    </div>""", unsafe_allow_html=True)

    col_l, col_m, col_r = st.columns([1,2,1])
    with col_m:
        st.markdown(f"### 🔐 {t('login')}")
        username = st.text_input(t("username"), placeholder=t("username_ph"))
        password = st.text_input(t("password"), type="password", placeholder=t("password_ph"))
        if st.button(t("login_btn"), type="primary", use_container_width=True):
            if username == st.secrets.get("APP_USERNAME","admin") and password == st.secrets.get("APP_PASSWORD","admin123"):
                st.session_state.logged_in = True
                st.rerun()
            else:
                st.error(t("login_err"))
    st.stop()

# ── SIDEBAR ───────────────────────────────────────────────────
with st.sidebar:
    st.markdown(f"### 📋 {t('title')}")
    st.markdown("---")

    # Dil & Tema
    lang_col, dark_col = st.columns([3, 2])
    with lang_col:
        lang = st.radio("🌐", ["TR","EN"], index=0 if st.session_state.lang=="TR" else 1,
                       horizontal=True, label_visibility="collapsed")
        if lang != st.session_state.lang:
            st.session_state.lang = lang
            st.rerun()
    with dark_col:
        mode_label = "☀️ Light" if st.session_state.dark else "🌙 Dark"
        if st.button(mode_label, key="theme_toggle", use_container_width=True):
            st.session_state.dark = not st.session_state.dark
            st.rerun()

    st.markdown("---")
    st.markdown(f"👤 **{st.secrets.get('APP_USERNAME','admin')}**")
    if st.button(t("start_over"), use_container_width=True):
        st.rerun()
    if st.button(t("logout"), use_container_width=True):
        st.session_state.logged_in = False
        st.rerun()
    st.markdown("---")
    st.markdown("<small style='color:#888'>Kapadokya University<br>🤖 DeepSeek AI</small>", unsafe_allow_html=True)

# ── HEADER ────────────────────────────────────────────────────
st.markdown(f"""<div class='app-header'>
    <h1>📋 {t('title')}</h1>
    <p>{t('subtitle')}</p>
</div>""", unsafe_allow_html=True)

# ── DEEPSEEK ──────────────────────────────────────────────────
DEEPSEEK_URL = "https://api.deepseek.com/chat/completions"

def deepseek_chat(messages, max_tokens=4096):
    resp = requests.post(
        DEEPSEEK_URL,
        headers={"Authorization": f"Bearer {st.secrets['DEEPSEEK_API_KEY']}", "Content-Type": "application/json"},
        json={"model": "deepseek-chat", "messages": messages, "max_tokens": max_tokens},
        timeout=60
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()

# ── SORU ÇIKARMA ──────────────────────────────────────────────
def extract_questions(text):
    sorular = []
    pattern1 = r'(\d{1,3}[\.\)]\s*)([\s\S]*?\?)'
    for idx, (_, soru) in enumerate(re.findall(pattern1, text), 1):
        soru = re.sub(r'\s+', ' ', soru.strip())
        if len(soru) > 8:
            sorular.append({"no": idx, "text": soru})
    if sorular:
        return sorular
    pattern2 = r'Soru\s+(\d+)\.\s+(.*?)(?=Soru\s+\d+\.|$)'
    for no_str, metin in re.findall(pattern2, text, re.DOTALL):
        metin = re.sub(r'\n\s*[A-E]\)\s.*', '', metin)
        metin = re.sub(r'Cevap:.*', '', metin, flags=re.DOTALL)
        metin = re.sub(r'\|', '', metin)
        metin = re.sub(r'\s+', ' ', metin.strip())
        if len(metin) > 8:
            sorular.append({"no": int(no_str), "text": metin})
    if sorular:
        return sorular
    for idx, line in enumerate(text.split('\n'), 1):
        line = line.strip()
        if line.endswith('?') and len(line) > 8:
            sorular.append({"no": idx, "text": line})
    return sorular

def extract_oc_from_pdf(pdf_bytes):
    import pypdf
    text = "\n".join(p.extract_text() or "" for p in pypdf.PdfReader(io.BytesIO(pdf_bytes)).pages)
    prompt = f"""Extract all learning outcomes (Dersin Öğrenme Çıktıları) from this Turkish university course info PDF.
Return ONLY JSON:
{{"outcomes":[{{"no":"LO-1","definition":"..."}},{{"no":"LO-2","definition":"..."}}]}}
TEXT: {text[:8000]}"""
    raw = deepseek_chat([{"role":"user","content":prompt}], max_tokens=2048)
    raw = re.sub(r'```json|```','',raw).strip()
    m = re.search(r'\{[\s\S]*\}', raw)
    if m:
        data = json.loads(m.group())
        return [{"no": o["no"], "tanim": o["definition"]} for o in data.get("outcomes", [])]
    return []

def auto_match(sorular, ocler):
    oc_listesi = "\n".join([f"- {o['no']}: {o['tanim']}" for o in ocler])
    result = {}
    n = len(sorular)
    chunks = (n + 19) // 20
    progress = st.progress(0, text=f"{t('mapping')} (1/{chunks})")

    for ci, i in enumerate(range(0, n, 20)):
        chunk = sorular[i:i+20]
        soru_listesi = "\n".join([f"{s['no']}. {s['text']}" for s in chunk])
        prompt = f"""You are an academic assessment expert. Match each exam question to learning outcomes AND classify by Bloom's Taxonomy.

LEARNING OUTCOMES:
{oc_listesi}

EXAM QUESTIONS:
{soru_listesi}

STRICT RULES:
- Match EACH question to EXACTLY ONE outcome (single match, 100%)
- ONLY assign multiple outcomes if the question EXPLICITLY covers 2+ distinct topics from different outcomes
- When in doubt, pick the SINGLE best matching outcome
- Percentages must sum to exactly 100 for each question

BLOOM'S TAXONOMY (pick exactly one):
- Remember: recall facts, definitions, lists (e.g. "What is X?", "Which method does Y?")
- Understand: explain, describe, summarize concepts (e.g. "Explain how X works")
- Apply: use knowledge in new situations, solve problems (e.g. "Which code does X?")
- Analyze: break down, compare, differentiate (e.g. "Compare X and Y", "Why does X cause Y?")
- Evaluate: judge, critique, justify decisions (e.g. "Which approach is better and why?")
- Create: design, build, formulate something new

Return ONLY valid JSON, no explanation:
{{"matches":[
  {{"q":1,"outcomes":[{{"lo":"LO-3","pct":100}}],"bloom":"Remember"}},
  {{"q":2,"outcomes":[{{"lo":"LO-1","pct":60}},{{"lo":"LO-2","pct":40}}],"bloom":"Understand"}}
]}}"""
        try:
            raw = deepseek_chat([{"role":"user","content":prompt}], max_tokens=2048)
            raw = re.sub(r'```json|```','',raw).strip()
            m = re.search(r'\{[\s\S]*\}', raw)
            if m:
                data = json.loads(m.group())
                for item in data.get("matches", []):
                    q_no = int(item.get("q", 0))
                    outcomes = item.get("outcomes", [])
                    bloom = str(item.get("bloom", "Remember"))
                    if q_no > 0 and outcomes:
                        result[q_no] = {"outcomes": outcomes, "zorluk": bloom}
        except Exception as e:
            st.warning(f"{t('batch')} {ci+1} error: {e}")
        progress.progress((ci+1)/chunks, text=f"{t('mapping')} ({ci+1}/{chunks})")

    progress.empty()
    return result

def build_rubrik(sorular, ocler, eslestirmeler, puan_esit, toplam_puan, ozel_puanlar,
                 ogretim_elemani="", bolum="", sinav_turu="", sinav_tarihi=""):
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

    KOYU_MOR = RGBColor(0x3B, 0x1F, 0x8C)
    MOR      = RGBColor(0x6D, 0x5C, 0xE7)
    BEYAZ    = RGBColor(0xFF, 0xFF, 0xFF)
    GRI_YAZ  = RGBColor(0x4A, 0x4A, 0x6A)

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def set_cell_border(cell, color="C4B5FD"):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   'single')
            border.set(qn('w:sz'),    '4')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def get_puan(no):
        if puan_esit:
            return round(toplam_puan / len(sorular), 1)
        return ozel_puanlar.get(no, round(toplam_puan / len(sorular), 1))

    oc_map = {o["no"]: o["tanim"] for o in ocler}

    doc = DocxDoc()

    # Sayfa marjinleri
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # ── BAŞLIK TABLOSU ─────────────────────────────────────
    t0 = doc.add_table(rows=1, cols=1)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    t0.style = 'Table Grid'
    c00 = t0.cell(0,0)
    set_cell_bg(c00, "3B1F8C")
    set_cell_border(c00)
    p0 = c00.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(8)
    p0.paragraph_format.space_after  = Pt(4)
    r0 = p0.add_run("SINAV RUBRİK FORMU")
    r0.bold      = True
    r0.font.size = Pt(16)
    r0.font.color.rgb = BEYAZ
    r0.font.name = "Arial"
    p0b = c00.add_paragraph("Kapadokya Üniversitesi")
    p0b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0b.paragraph_format.space_after = Pt(6)
    rb = p0b.add_run()
    rb.font.size  = Pt(10)
    rb.font.color.rgb = RGBColor(0xC4, 0xB5, 0xFD)
    rb.font.name = "Arial"
    # fix: remove empty run
    p0b.runs[0].text = "Kapadokya Üniversitesi"

    doc.add_paragraph()

    # ── BİLGİ TABLOSU ──────────────────────────────────────
    t1 = doc.add_table(rows=3, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.style = 'Table Grid'

    info = [
        [("Öğretim Elemanı:", ogretim_elemani or "—"), ("Bölüm/Program:", bolum or "—")],
        [("Dersin Adı:", exam_filename if 'exam_filename' in dir() else "—"), ("Sınav Türü:", sinav_turu or "—")],
        [("Sınav Tarihi:", sinav_tarihi or "—"), ("Doküman No:", "DKM.FR.033")],
    ]

    for ri, row in enumerate(info):
        bg = "EDE9FE" if ri % 2 == 0 else "FFFFFF"
        for ci, (label, value) in enumerate(row):
            cell = t1.cell(ri, ci)
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            rl = p.add_run(label + " ")
            rl.bold = True
            rl.font.size = Pt(10)
            rl.font.color.rgb = KOYU_MOR
            rl.font.name = "Arial"
            rv = p.add_run(value)
            rv.font.size = Pt(10)
            rv.font.color.rgb = GRI_YAZ
            rv.font.name = "Arial"

    doc.add_paragraph()

    # ── ÖÇ LİSTESİ ─────────────────────────────────────────
    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(4)
    ph.paragraph_format.space_after  = Pt(4)
    rh = ph.add_run("Dersin Öğrenme Çıktıları")
    rh.bold = True
    rh.font.size = Pt(11)
    rh.font.color.rgb = KOYU_MOR
    rh.font.name = "Arial"

    for oc in ocler:
        po = doc.add_paragraph(style='List Number')
        po.paragraph_format.space_before = Pt(2)
        po.paragraph_format.space_after  = Pt(2)
        ro = po.add_run(oc["tanim"])
        ro.font.size = Pt(10)
        ro.font.color.rgb = GRI_YAZ
        ro.font.name = "Arial"

    doc.add_paragraph()

    # ── TABLO BAŞLIĞI ───────────────────────────────────────
    pt = doc.add_paragraph()
    pt.paragraph_format.space_before = Pt(6)
    pt.paragraph_format.space_after  = Pt(4)
    rt = pt.add_run("Tablo-1. Çoktan Seçmeli Sorular İçin Dereceli Puanlama Anahtarı")
    rt.bold = True
    rt.font.size = Pt(11)
    rt.font.color.rgb = KOYU_MOR
    rt.font.name = "Arial"

    # ── ANA TABLO ───────────────────────────────────────────
    headers = ["Soru No", "Puan", "Doğru", "Yanlış", "Ölçülen Kazanım", "Ders Öğrenme Çıktısı", "Bloom Taksonomisi"]
    col_widths = [Cm(1.3), Cm(1.2), Cm(1.3), Cm(1.3), Cm(6.5), Cm(3.5), Cm(2.5)]

    t2 = doc.add_table(rows=1+len(sorular)+1, cols=len(headers))
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.style = 'Table Grid'

    # Header satırı
    for ci, (htext, cw) in enumerate(zip(headers, col_widths)):
        cell = t2.cell(0, ci)
        cell.width = cw
        set_cell_bg(cell, "3B1F8C")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = BEYAZ
        r.font.name = "Arial"

    # Soru satırları
    for ri, s in enumerate(sorular):
        esl      = eslestirmeler.get(s["no"], {})
        outcomes = esl.get("outcomes", [])
        bloom    = esl.get("zorluk", "")
        oc_str   = ", ".join(f"{o['lo']}(%{o['pct']})" for o in outcomes) if outcomes else "—"
        puan     = get_puan(s["no"])
        bg       = "F5F4FF" if ri % 2 == 0 else "FFFFFF"

        row_data = [
            (str(s["no"]), WD_ALIGN_PARAGRAPH.CENTER),
            (str(puan),    WD_ALIGN_PARAGRAPH.CENTER),
            ("4",          WD_ALIGN_PARAGRAPH.CENTER),
            ("0",          WD_ALIGN_PARAGRAPH.CENTER),
            (s["text"],    WD_ALIGN_PARAGRAPH.LEFT),
            (oc_str,       WD_ALIGN_PARAGRAPH.CENTER),
            (bloom,        WD_ALIGN_PARAGRAPH.CENTER),
        ]

        for ci, (txt, align) in enumerate(row_data):
            cell = t2.cell(ri+1, ci)
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = align
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(txt)
            r.font.size = Pt(9)
            r.font.color.rgb = GRI_YAZ
            r.font.name = "Arial"

    # Toplam satırı
    last_row = len(sorular) + 1
    for ci in range(len(headers)):
        cell = t2.cell(last_row, ci)
        set_cell_bg(cell, "3B1F8C")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        text = "TOPLAM" if ci == 0 else (str(toplam_puan) if ci == 1 else "0" if ci == 3 else "")
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = BEYAZ
        r.font.name = "Arial"

    # ── DİPNOT ─────────────────────────────────────────────
    pd = doc.add_paragraph()
    pd.paragraph_format.space_before = Pt(8)
    rd = pd.add_run("* Yanıt biçimi istenildiği takdirde alt detaylara bölünerek çoğaltılabilir.")
    rd.italic = True
    rd.font.size = Pt(9)
    rd.font.color.rgb = GRI_YAZ
    rd.font.name = "Arial"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_excel(sorular, ocler, eslestirmeler, anahtar, puan_esit, toplam_puan, ozel_puanlar):
    wb = Workbook()
    PURPLE="5B4FCF"; TEAL="0F766E"; GREEN_L="F0FDF4"
    GRAY="F8F7FF"; WHITE="FFFFFF"; DARK="1A1640"

    def hstyle(cell, bg=PURPLE):
        cell.font = Font(bold=True, color=WHITE, size=10)
        cell.fill = PatternFill("solid", fgColor=bg)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def add_borders(ws, r1, r2, c1, c2):
        t = Side(style="thin", color="E2E0F0")
        for row in ws.iter_rows(min_row=r1, max_row=r2, min_col=c1, max_col=c2):
            for cell in row:
                cell.border = Border(left=t, right=t, top=t, bottom=t)

    oc_no_to_idx = {o["no"]: i+1 for i, o in enumerate(ocler)}

    def get_puan(soru_no):
        if puan_esit:
            return round(toplam_puan / len(sorular), 2)
        return ozel_puanlar.get(soru_no, round(toplam_puan / len(sorular), 2))

    max_oc = max((len(eslestirmeler.get(s["no"],{}).get("outcomes",[])) for s in sorular), default=1)
    max_oc = max(max_oc, 1)

    # ─── SAYFA 1: Question-LO Mapping ───────────────────────
    ws1 = wb.active
    ws1.title = "Question-LO Mapping"
    ws1.row_dimensions[1].height = 36

    headers1 = ["Q NO", "SCORE", "QUESTION"]
    for i in range(max_oc):
        headers1.append(f"DÇ Sıra {i+1}")
        headers1.append(f"Etki Oran {i+1}")
    headers1 += ["BLOOM TAKSONOMİSİ", "ANSWER KEY"]

    for col, h in enumerate(headers1, 1):
        hstyle(ws1.cell(row=1, column=col, value=h))

    for ri, s in enumerate(sorular, 2):
        esl = eslestirmeler.get(s["no"], {})
        outcomes = esl.get("outcomes", [])
        diff = esl.get("zorluk","")
        key = anahtar[s["no"]-1] if anahtar and s["no"]-1 < len(anahtar) else "-"
        puan = get_puan(s["no"])

        ws1.cell(ri,1,s["no"]).alignment = Alignment(horizontal="center")
        ws1.cell(ri,2,puan).alignment    = Alignment(horizontal="center")
        ws1.cell(ri,3,s["text"]).alignment = Alignment(wrap_text=True, vertical="top")

        for i, outcome in enumerate(outcomes):
            lo_idx = oc_no_to_idx.get(outcome.get("lo",""), "")
            pct    = outcome.get("pct", 100)
            cb = 4 + i*2
            ws1.cell(ri, cb,   lo_idx).alignment = Alignment(horizontal="center")
            ws1.cell(ri, cb+1, pct).alignment    = Alignment(horizontal="center")

        diff_col = 4 + max_oc*2
        ws1.cell(ri, diff_col,   diff).alignment = Alignment(horizontal="center")
        ws1.cell(ri, diff_col+1, key).alignment  = Alignment(horizontal="center")

        if outcomes:
            for c in range(1, len(headers1)+1):
                ws1.cell(ri,c).fill = PatternFill("solid", fgColor="EDE9FE")
        elif ri%2==0:
            for c in range(1, len(headers1)+1):
                ws1.cell(ri,c).fill = PatternFill("solid", fgColor=GRAY)

    ws1.column_dimensions["A"].width = 8
    ws1.column_dimensions["B"].width = 8
    ws1.column_dimensions["C"].width = 58
    for i in range(max_oc):
        ws1.column_dimensions[get_column_letter(4+i*2)].width   = 10
        ws1.column_dimensions[get_column_letter(4+i*2+1)].width = 12
    ws1.column_dimensions[get_column_letter(4+max_oc*2)].width   = 12
    ws1.column_dimensions[get_column_letter(4+max_oc*2+1)].width = 12
    ws1.freeze_panes = "A2"
    add_borders(ws1, 1, len(sorular)+1, 1, len(headers1))

    # ─── SAYFA 2: LO Summary ────────────────────────────────
    ws2 = wb.create_sheet("LO Summary")
    ws2.row_dimensions[1].height = 36
    for col, h in enumerate(["LO NO","LO DEFINITION","# QUESTIONS","QUESTION NUMBERS"],1):
        hstyle(ws2.cell(row=1, column=col, value=h), bg=TEAL)

    for ri, oc in enumerate(ocler, 2):
        matched = [s for s in sorular if any(o["lo"]==oc["no"] for o in eslestirmeler.get(s["no"],{}).get("outcomes",[]))]
        q_nums = ", ".join(str(s["no"]) for s in matched)
        ws2.cell(ri,1,oc["no"]).alignment = Alignment(horizontal="center")
        ws2.cell(ri,2,oc["tanim"]).alignment = Alignment(wrap_text=True)
        ws2.cell(ri,3,len(matched)).alignment = Alignment(horizontal="center")
        ws2.cell(ri,4,q_nums).alignment = Alignment(wrap_text=True)
        if ri%2==0:
            for c in [1,2,3,4]: ws2.cell(ri,c).fill = PatternFill("solid",fgColor="F0FDFA")

    for col,w in zip("ABCD",[10,52,14,40]):
        ws2.column_dimensions[col].width = w
    ws2.freeze_panes = "A2"
    add_borders(ws2, 1, len(ocler)+1, 1, 4)

    # ─── SAYFA 3: Proliz Format ─────────────────────────────
    ws3 = wb.create_sheet("Proliz Format")
    ws3.row_dimensions[1].height = 36

    headers3 = ["Soru No", "Soru Puan", "Soru Metni"]
    for i in range(max_oc):
        headers3.append(f"DÇ Sıra {i+1}")
        headers3.append(f"Etki Oran {i+1}")
    headers3.append("Bloom Taksonomisi")

    for col, h in enumerate(headers3, 1):
        hstyle(ws3.cell(row=1, column=col, value=h))

    for ri, s in enumerate(sorular, 2):
        esl = eslestirmeler.get(s["no"], {})
        outcomes = esl.get("outcomes", [])
        puan = get_puan(s["no"])

        ws3.cell(ri,1,s["no"]).alignment = Alignment(horizontal="center")
        ws3.cell(ri,2,puan).alignment    = Alignment(horizontal="center")
        ws3.cell(ri,3,s["text"]).alignment = Alignment(wrap_text=True, vertical="top")

        for i, outcome in enumerate(outcomes):
            lo_idx = oc_no_to_idx.get(outcome.get("lo",""), "")
            pct    = outcome.get("pct", 100)
            cb = 4 + i*2
            ws3.cell(ri, cb,   lo_idx).alignment = Alignment(horizontal="center")
            ws3.cell(ri, cb+1, pct).alignment    = Alignment(horizontal="center")

        # Bloom sütunu
        bloom_col = 4 + max_oc*2 + 1
        bloom_val = esl.get("zorluk", "")
        ws3.cell(ri, bloom_col, bloom_val).alignment = Alignment(horizontal="center")

        if outcomes:
            for c in range(1, len(headers3)+1):
                ws3.cell(ri,c).fill = PatternFill("solid", fgColor="EDE9FE")
        elif ri%2==0:
            for c in range(1, len(headers3)+1):
                ws3.cell(ri,c).fill = PatternFill("solid", fgColor=GRAY)

    ws3.column_dimensions["A"].width = 9
    ws3.column_dimensions["B"].width = 10
    ws3.column_dimensions["C"].width = 60
    for i in range(max_oc):
        ws3.column_dimensions[get_column_letter(4+i*2)].width   = 10
        ws3.column_dimensions[get_column_letter(4+i*2+1)].width = 12
    ws3.freeze_panes = "A2"
    add_borders(ws3, 1, len(sorular)+1, 1, len(headers3))

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════
# UI
# ══════════════════════════════════════════════════════════════

# ADIM 1
st.markdown(f"<div class='section-title'>{t('step1')}</div>", unsafe_allow_html=True)
uploaded_exam = st.file_uploader(t("step1_sub"), type=["txt","pdf","docx","doc"], key="exam", label_visibility="collapsed")

sorular = []
exam_filename = "Exam-Outcome-Report"
if uploaded_exam:
    exam_filename = uploaded_exam.name.rsplit(".",1)[0]
    ext = uploaded_exam.name.split(".")[-1].lower()
    raw = uploaded_exam.read()
    try:
        if ext == "txt":
            text = raw.decode("utf-8", errors="ignore")
            sorular = extract_questions(text)
        elif ext in ["docx","doc"]:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            para_text = "\n".join(p.text for p in doc.paragraphs)
            sorular = extract_questions(para_text)
            if not sorular:
                seen = set()
                table_text = ""
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            ct = cell.text.strip()
                            if ct and ct not in seen:
                                seen.add(ct)
                                table_text += ct + "\n"
                sorular = extract_questions(table_text)
        elif ext == "pdf":
            import pypdf
            text = "\n".join(p.extract_text() or "" for p in pypdf.PdfReader(io.BytesIO(raw)).pages)
            sorular = extract_questions(text)

        if sorular:
            st.success(f"✅ {len(sorular)} {t('step1_ok')}")
            with st.expander(f"👁 {t('preview')}", expanded=False):
                for s in sorular:
                    st.markdown(f"**{s['no']}.** {s['text'][:120]}...")
        else:
            st.error(t("step1_err"))
    except Exception as e:
        st.error(f"❌ {e}")

st.divider()

# ADIM 2
st.markdown(f"<div class='section-title'>{t('step2')}</div>", unsafe_allow_html=True)
tab1, tab2 = st.tabs([t("tab_pdf"), t("tab_manual")])

ocler = []
with tab1:
    st.info(t("pdf_info"))
    uploaded_bp = st.file_uploader(t("pdf_upload"), type=["pdf"], key="bilgi")
    if uploaded_bp:
        with st.spinner(t("oc_extracting")):
            try:
                extracted = extract_oc_from_pdf(uploaded_bp.read())
                if extracted:
                    st.success(f"✅ {len(extracted)} {t('pdf_ok')}")
                    for oc in extracted:
                        st.markdown(f"**{oc['no']}** — {oc['tanim'][:80]}...")
                    ocler = extracted
                else:
                    st.error(t("pdf_err"))
            except Exception as e:
                st.error(f"❌ {e}")

with tab2:
    st.info(t("manual_info"))
    oc_manual = st.text_area("", height=160, placeholder=t("manual_ph"))
    if oc_manual.strip():
        for line in oc_manual.strip().split("\n"):
            if ":" in line:
                no, tanim = line.split(":",1)
                no, tanim = no.strip(), tanim.strip()
                if no and tanim:
                    ocler.append({"no":no,"tanim":tanim})
        if ocler:
            st.success(f"✅ {len(ocler)} {t('outcomes_ok')}")

st.divider()

# ADIM 3
st.markdown(f"<div class='section-title'>{t('step3')}</div>", unsafe_allow_html=True)
col1, col2 = st.columns(2)
with col1:
    toplam_puan = st.number_input(t("total_pts"), min_value=1, max_value=1000, value=100)
with col2:
    puan_tipi = st.radio(t("scoring_type"), [t("equal"), t("custom")], horizontal=True)

puan_esit = puan_tipi == t("equal")
ozel_puanlar = {}

if not puan_esit and sorular:
    st.markdown(f"**{t('enter_scores')}**")
    cols = st.columns(5)
    for i, s in enumerate(sorular):
        with cols[i % 5]:
            ozel_puanlar[s["no"]] = st.number_input(
                f"Q{s['no']}", min_value=0.0, max_value=float(toplam_puan),
                value=round(toplam_puan/len(sorular),1),
                key=f"puan_{s['no']}"
            )
    total_check = sum(ozel_puanlar.values())
    if abs(total_check - toplam_puan) > 0.5:
        st.warning(f"{t('total_warn')} {total_check} (= {toplam_puan})")
    else:
        st.success(f"{t('total_ok')} {total_check}")

st.divider()

# ADIM 4: Cevap anahtarı
with st.expander(t("answer_key")):
    anahtar = st.text_input("", placeholder=t("answer_ph")).upper()

st.divider()

# RUBRİK BİLGİLERİ
st.markdown(f"<div class='section-title'>📋 {t('rubrik_info')}</div>", unsafe_allow_html=True)
rcol1, rcol2 = st.columns(2)
with rcol1:
    ogretim_elemani = st.text_input(t("ogretim_elemani"), placeholder="Öğr. Gör. Ad Soyad")
    sinav_turu      = st.text_input(t("sinav_turu"), placeholder="Vize / Final / Quiz")
with rcol2:
    bolum_rubrik   = st.text_input(t("bolum"), placeholder="Bilgisayar Programcılığı")
    sinav_tarihi   = st.text_input(t("sinav_tarihi"), placeholder="01/01/2026")

st.divider()
st.markdown(f"<div class='section-title'>{t('step4')}</div>", unsafe_allow_html=True)

ready = bool(sorular) and bool(ocler)
if not ready:
    if not sorular: st.warning(t("no_exam"))
    if not ocler:   st.warning(t("no_oc"))
else:
    if puan_esit:
        puan_per_q = round(toplam_puan / len(sorular), 2)
        st.success(f"{t('ready_equal')} **{len(sorular)} {t('questions')}** × **{len(ocler)} {t('outcomes')}** — {puan_per_q} {t('ready_pts')}")
    else:
        st.success(f"{t('ready_equal')} **{len(sorular)} {t('questions')}** × **{len(ocler)} {t('outcomes')}** — {t('ready_custom')}")

    if st.button(t("map_btn"), type="primary", use_container_width=True):
        try:
            with st.spinner(t("mapping")):
                eslestirmeler = auto_match(sorular, ocler)
            mapped = sum(1 for v in eslestirmeler.values() if v.get("outcomes"))
            multi  = sum(1 for v in eslestirmeler.values() if len(v.get("outcomes",[])) > 1)

            # Session state'e kaydet — butonlar kaybolmasın
            st.session_state["son_eslestirmeler"]  = eslestirmeler
            st.session_state["son_sorular"]        = sorular
            st.session_state["son_ocler"]          = ocler
            st.session_state["son_anahtar"]        = anahtar
            st.session_state["son_puan_esit"]      = puan_esit
            st.session_state["son_toplam_puan"]    = toplam_puan
            st.session_state["son_ozel_puanlar"]   = ozel_puanlar
            st.session_state["son_exam_filename"]  = exam_filename
            st.session_state["son_ogretim"]        = ogretim_elemani
            st.session_state["son_bolum"]          = bolum_rubrik
            st.session_state["son_sinav_turu"]     = sinav_turu
            st.session_state["son_sinav_tarihi"]   = sinav_tarihi
            st.session_state["son_mapped"]         = mapped
            st.session_state["son_multi"]          = multi
            st.rerun()
        except Exception as e:
            st.error(f"❌ {e}")

    # Sonuçlar session state'te varsa her zaman göster
    if "son_eslestirmeler" in st.session_state and st.session_state["son_sorular"] == sorular:
        esl_  = st.session_state["son_eslestirmeler"]
        sor_  = st.session_state["son_sorular"]
        oc_   = st.session_state["son_ocler"]
        fn_   = st.session_state["son_exam_filename"]

        st.success(f"✅ {st.session_state['son_mapped']}/{len(sor_)} {t('mapped_ok')} — {st.session_state['son_multi']} {t('multi_match')}")

        col1, col2 = st.columns(2)
        with col1:
            excel_bytes = build_excel(
                sor_, oc_, esl_,
                st.session_state["son_anahtar"],
                st.session_state["son_puan_esit"],
                st.session_state["son_toplam_puan"],
                st.session_state["son_ozel_puanlar"]
            )
            st.download_button(
                t("download"), excel_bytes,
                f"{fn_}-Outcome-Report.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True, type="primary",
                key="dl_excel"
            )
        with col2:
            rubrik_bytes = build_rubrik(
                sor_, oc_, esl_,
                st.session_state["son_puan_esit"],
                st.session_state["son_toplam_puan"],
                st.session_state["son_ozel_puanlar"],
                st.session_state["son_ogretim"],
                st.session_state["son_bolum"],
                st.session_state["son_sinav_turu"],
                st.session_state["son_sinav_tarihi"]
            )
            st.download_button(
                "📄 " + t("download_rubrik"), rubrik_bytes,
                f"{fn_}-Rubrik.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_rubrik"
            )
